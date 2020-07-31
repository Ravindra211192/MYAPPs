#!/home/developer/Desktop/openHTF/python_virenv/bin/python3

from __future__ import print_function
import openhtf as htf
import os
import shutil
import tempfile
import threading
import queue
import time
import locale, datetime
import pexpect
import paramiko
import logging
import collections
import config_parser
import netifaces as ni 
from openpyxl import Workbook
from openpyxl.styles import Font, Color, Alignment, Border, Side, NamedStyle, colors 
from scp import SCPClient 
from pexpect import pxssh
from openhtf.util import conf
from openhtf.output.servers import station_server
from openhtf.output.web_gui import web_launcher
from openhtf.plugs import user_input
from openhtf import PhaseOptions
from openhtf import plugs
from git import Repo
from docx import Document
from docx.shared import RGBColor,Pt,Inches
from docx.text.parfmt import ParagraphFormat
from docx.enum.text import WD_ALIGN_PARAGRAPH

'''------------------------------------------------------------------------------------------------------------------------------------------------------------------

                  Class FS_BoardPlug(plugs.BasePlug): class Plug to establish connection to the boards and thereafter test the interfaces according to the user demand

                  Input parameters:

                  plugs.BasePlug           --> Instance of a class plug FS_BoardPlug

  ------------------------------------------------------------------------------------------------------------------------------------------------------------------------'''    

'''Plug to access the boards'''   
class FS_BoardPlug(plugs.BasePlug):

  """This plug simply does a ssh connection to the host attribute."""

  dut_id_test_board='0.0.0.0'
  dut_id_testing_board='0.0.0.0'
  name=None
  sel_interface=None
  ret_code_s=0
  ret_code_r=0
  enable_remote=True
  username='root'
  password='root'
  branch='master'
  depth =1
  
  global select_interface
  global select_commandline
  global select_board
  global test_phases
  global test

  def __init__(self):

    FS_BoardPlug.enable_remote=True
    FS_BoardPlug.auto_prompt_reset=True
    FS_BoardPlug.ssh_key=True
    assert self.dut_id is not None

  '''------------------------------------------------------------------------------------------------------------------------------------------------------------------
                  Class method pxssh_login(self): establishes a SSH connection to the boards using openhtf utility
                  Input parameters:
                  self           --> Instance of a class plug FS_BoardPlug
                  Tipp: If in case an application couldnot establish a ssh connection to the board, then clear the whole contents from the known_hosts file, which is located 
                  under /home/developer/.ssh/ directory.
  ------------------------------------------------------------------------------------------------------------------------------------------------------------------------'''    

  def pxssh_login(self):    
    
        @PhaseOptions(timeout_s=60*60*24)
        @plugs.plug(prompts=user_input.UserInput)
        def start_phase(test, prompts):
             
          """Name of a tester"""
          FS_BoardPlug.name = prompts.prompt(message='Name of a tester',text_input=True,
                                timeout_s=60*60*24)
          #test.test_record.dut_id = FS_BoardPlug.name
          print('Testers’ name \n')
          """Login for the another board next to the to be tested board"""          
          FS_BoardPlug.dut_id_testing_board = prompts.prompt(message='Enter an IP address of the Partner board(efusA9).',text_input=True,
                                timeout_s=60*60*24)
          test.test_record.dut_id = FS_BoardPlug.dut_id_testing_board
          self.px_ssh_1=pxssh.pxssh()
          self.px_ssh_1.login(FS_BoardPlug.dut_id_testing_board,FS_BoardPlug.username,FS_BoardPlug.password)
          print('Connection to the testing board(efusa9) via SSH is successfull!\n') 
         
          """Test start trigger that prompts the user for a DUT ID."""
          FS_BoardPlug.dut_id_test_board = prompts.prompt(message='Enter an IP address of the board to be tested.',text_input=True,
                                timeout_s=60*60*24)          
          print("DUT:",FS_BoardPlug.dut_id_test_board)
          test.test_record.dut_id = FS_BoardPlug.dut_id_test_board
          self.px_ssh=pxssh.pxssh()
          self.px_ssh.login(FS_BoardPlug.dut_id_test_board,FS_BoardPlug.username,FS_BoardPlug.password) 
          print('Connection to the test board via SSH is successfull!\n')
                      
        return start_phase
       
  '''------------------------------------------------------------------------------------------------------------------------------------------------------------------
                  Class method get_exe_file(self,interface): extracting files for the interface testing purpose from Gitlab repo
                  Input parameters:
                  self           --> Instance of a class plug FS_BoardPlug
                  interface      --> name of an interface so that it can be used to identify and transfer the file  
  ------------------------------------------------------------------------------------------------------------------------------------------------------------------------'''                   
  '''extracting files from Gitlab repo'''
  def get_exe_file(self,interface):
   
    self.interface=interface
    self.temp=tempfile.mkdtemp()
    
    os.system('git config --global credential.helper store')
    os.system('echo "https://RC_FS:WjQm77ah@gitlab.com" > ~/.git-credentials')
    os.system('git config --global user.name "RC_FS"')
    os.system('git config --global user.email "chand@fs-net.de"')

    Repo.clone_from('https://gitlab.com/RC_FS/interface-testing.git', self.temp)
    
    if os.path.isfile('/home/developer/Desktop/selftest_i2c'):
      print("File selftest_i2c exists!")
    elif(self.interface=='i2c'):   
      shutil.move(os.path.join(self.temp,'selftest_i2c'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/selftest_spi'):
      print("File selftest_spi exists!")
    elif(self.interface=='spi'):  
      shutil.move(os.path.join(self.temp,'selftest_spi'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/send_uart_abr'):
      print("File send_uart_abr exists!")
    elif(self.interface=='send_uart_abr'): 
      shutil.move(os.path.join(self.temp,'send_uart_abr'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/send_uart_rs485_abr'):
      print("File send_uart_rs485_abr exists!")
    elif(self.interface=='send_uart_rs485_abr'): 
      shutil.move(os.path.join(self.temp,'send_uart_rs485_abr'), '/home/developer/Desktop/')
     
    if os.path.isfile('/home/developer/Desktop/rcv_uart_abr'):
      print("File rcv_uart_abr exists!")
    elif(self.interface=='rcv_uart_abr'): 
      shutil.move(os.path.join(self.temp,'rcv_uart_abr'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/rcv_uart_rs485_abr'):
      print("File rcv_uart_rs485_abr exists!")
    elif(self.interface=='rcv_uart_rs485_abr'): 
      shutil.move(os.path.join(self.temp,'rcv_uart_rs485_abr'), '/home/developer/Desktop/')
    
    if os.path.isfile('/home/developer/Desktop/snd_can_afs'):
      print("File snd_can_afs exists!")
    elif(self.interface=='snd_can_afs'):  
      shutil.move(os.path.join(self.temp,'snd_can_afs'), '/home/developer/Desktop/')
     
    if os.path.isfile('/home/developer/Desktop/rcv_can_afs'):
      print("File rcv_can_afs exists!")
    elif(self.interface=='rcv_can_afs'):
      shutil.move(os.path.join(self.temp,'rcv_can_afs'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/selftest_pwm_new'):
      print("File selftest_pwm_new exists!")
    elif(self.interface=='pwm'):
      shutil.move(os.path.join(self.temp,'selftest_pwm_new'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/selftest_network'):
      print("File selftest_network exists!")
    elif(self.interface=='network'):
      shutil.move(os.path.join(self.temp,'selftest_network'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/selftest_sdcard'):
      print("File selftest_sdcard exists!")
    elif(self.interface=='sdcard'):
      shutil.move(os.path.join(self.temp,'selftest_sdcard'), '/home/developer/Desktop/')
      
    if os.path.isfile('/home/developer/Desktop/selftest_flash.sh'):
      print("File selftest_flash.sh exists!")
    elif(self.interface=='flash'):
      shutil.move(os.path.join(self.temp,'selftest_flash.sh'), '/home/developer/Desktop/')
      
    shutil.rmtree(self.temp)
    
    pass
    
  '''------------------------------------------------------------------------------------------------------------------------------------------------------------------
                  Class method send_file_to_remote(self,select_file): sending the executable files to the remote boards
                  Input parameters:
                  self           --> Instance of a class plug FS_BoardPlug
                  select_file    --> name of an interface so that it identifies the correct file from the downloaded files and send it to the boards using SCP utility 
  ------------------------------------------------------------------------------------------------------------------------------------------------------------------------'''                     

  def send_file_to_remote(self,select_file):

    self.select_file=select_file    
    self.ssh=paramiko.SSHClient()
    self.ssh.load_system_host_keys()
    self.ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    self.ssh.connect(FS_BoardPlug.dut_id_test_board, 22, "root", "root")
    self.scp =SCPClient(self.ssh.get_transport())

    self.ssh_1=paramiko.SSHClient()
    self.ssh_1.load_system_host_keys()
    self.ssh_1.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    self.ssh_1.connect(FS_BoardPlug.dut_id_testing_board, 22, "root", "root")
    self.scp_1 =SCPClient(self.ssh_1.get_transport())

    if select_file=='i2c':
    
      self.scp.put('/home/developer/Desktop/selftest_i2c','/root/')
      self.scp_1.put('/home/developer/Desktop/selftest_i2c','/root/')
      print("selftest_i2c copied successfully!")

    elif select_file=='spi':
    
      self.scp.put('/home/developer/Desktop/selftest_spi','/root/')
      self.scp_1.put('/home/developer/Desktop/selftest_spi','/root/')
      print("selftest_spi copied successfully!")

    elif select_file=='uart':
      self.scp.put('/home/developer/Desktop/rcv_uart_abr','/root/')
      self.scp.put('/home/developer/Desktop/rcv_uart_rs485_abr','/root/') 
      self.scp_1.put('/home/developer/Desktop/rcv_uart_abr','/root/')
      self.scp_1.put('/home/developer/Desktop/rcv_uart_rs485_abr','/root/')  
      print("rcv_uart_abr copied successfully!")
      print("rcv_uart_rs485_abr copied successfully!")
      self.scp.put('/home/developer/Desktop/send_uart_abr','/root/')
      self.scp.put('/home/developer/Desktop/send_uart_rs485_abr','/root/') 
      self.scp_1.put('/home/developer/Desktop/send_uart_abr','/root/')
      self.scp_1.put('/home/developer/Desktop/send_uart_rs485_abr','/root/')
      print("send_uart_rs485_abr copied successfully!")
      print("send_uart_abr copied successfully!")

    elif select_file=='can':

      self.scp.put('/home/developer/Desktop/rcv_can_afs','/root/')
      self.scp_1.put('/home/developer/Desktop/rcv_can_afs','/root/')  
      print("rcv_can_afs copied successfully!")
      self.scp.put('/home/developer/Desktop/snd_can_afs','/root/') 
      self.scp_1.put('/home/developer/Desktop/snd_can_afs','/root/') 
      print("snd_can_afs copied successfully!")

    elif select_file=='pwm':
    
      self.scp.put('/home/developer/Desktop/selftest_pwm_new','/root/')
      self.scp_1.put('/home/developer/Desktop/selftest_pwm_new','/root/')
      print("selftest_pwm_new copied successfully!")
      
    elif select_file=='network':

      self.scp.put('/home/developer/Desktop/selftest_network','/root/')
      self.scp_1.put('/home/developer/Desktop/selftest_network','/root/')
      print("selftest_network copied successfully!")
      
    elif select_file=='sdcard':

      self.scp.put('/home/developer/Desktop/selftest_sdcard','/root/')
      self.scp_1.put('/home/developer/Desktop/selftest_sdcard','/root/')
      print("selftest_sdcard copied successfully!")
      
    elif select_file=='flash':

      self.scp.put('/home/developer/Desktop/selftest_flash.sh','/root/')
      self.scp_1.put('/home/developer/Desktop/selftest_flash.sh','/root/')
      print("selftest_flash.sh copied successfully!")

    self.scp.close() 
    self.scp_1.close()
  '''------------------------------------------------------------------------------------------------------------------------------------------------------------------
                  class method print_console(self,ssh_select): prints out the whole console
                  Input parameters:
                  self           --> Instance of a class plug FS_BoardPlug
                  ssh_select     --> select ssh prompts (two prompts is created beacause of two boards)  ------------------------------------------------------------------------------------------------------------------------------------------------------------------------'''  
  def print_console(self,ssh_select):
  
    if ssh_select==0:
      self.px_ssh.prompt()
      print(self.px_ssh.before)    

    else:
      self.px_ssh_1.prompt()
      print(self.px_ssh_1.before)
  
  def mount_rw(self, board):
  
    self.board=board

    if self.board=='efusa9':
      self.px_ssh.sendline('mount -o remount, rw /')
      self.print_console(self,0)      

    elif self.board=='testboard':
      self.px_ssh_1.sendline('mount -o remount, rw /')
      self.print_console(self,1)
      
  def snd_return_code_call(self,q_s,board):
      
      self.board=board      
      if self.board=='pxssh':
        self.px_ssh.sendline('echo $?')
        self.print_console(self,0) 
        ###Return code stores at the index 9, i.e. self.ret_code_s_can[9]  
        self.ret_code_s_can= '\r\n'.join(str(self.px_ssh.before.splitlines()[1:]))

      else:
        self.px_ssh_1.sendline('echo $?')
        self.print_console(self,1)   
        self.ret_code_s_can= '\r\n'.join(str(self.px_ssh_1.before.splitlines()[1:]))
        
      q_s.put(self.ret_code_s_can[9])    
   
  def rcv_return_code_call(self,q_r,board):
      
      self.board=board
      if self.board=='pxssh':
        self.px_ssh.sendline('echo $?')
        self.print_console(self,0)  
        self.ret_code_r_can= '\r\n'.join(str(self.px_ssh.before.splitlines()[1:]))

      else:

        self.px_ssh_1.sendline('echo $?')
        self.print_console(self,1)  
        self.ret_code_r_can= '\r\n'.join(str(self.px_ssh_1.before.splitlines()[1:]))     
      q_r.put(self.ret_code_r_can[9])
             
  def uart_referenceboard(self, sel_snd,q_s,q_r,uart):
      
    self.sel_snd= sel_snd
    self.q_s= q_s
    self.q_r= q_r
    self.uart= uart
    
    if self.uart in config_parser.interfacenamelistref:
      indx=config_parser.interfacenamelistref.index(self.uart)
      if config_parser.interfacenamelistref[indx+1].find(",") >=0:
        self.type_list=config_parser.interfacenamelistref[indx+1].split(",")   
        self.send_commandline=self.type_list[0]        
        self.rcv_commandline=self.type_list[1]
        print("Send commandline: {}".format(self.send_commandline))
        print("Receive commandline: {}".format(self.rcv_commandline))  
                  
    if self.sel_snd=='send':
      
      self.px_ssh_1.sendline(str(self.send_commandline))  
      self.print_console(self,1)     
      self.snd_return_code_call(self,self.q_s,'pxssh1')

    else:

      self.px_ssh_1.sendline(str(self.rcv_commandline))
      self.print_console(self,1)
      self.rcv_return_code_call(self,self.q_r,'pxssh1') 
      
  def testingboard(self, sel_rcv,q_s,q_r,commandline):          
    self.sel_rcv=sel_rcv
    self.q_s=q_s
    self.q_r=q_r

    if commandline.find(",") >= 0:
          self.type_list=commandline.split(",")
          self.rcv_commandline=self.type_list[0]
          self.rcv_commandline=self.rcv_commandline.split("*")
          self.rcv_commandline=self.rcv_commandline[0]
          self.send_commandline=self.type_list[1]
          self.send_commandline=self.send_commandline.split("*")
          self.send_commandline=self.send_commandline[0]
    else:
          pass  
    if self.sel_rcv== 'rcv':
      self.px_ssh.sendline(str(self.rcv_commandline))
      self.print_console(self,0)
      self.rcv_return_code_call(self,self.q_r,'pxssh') 
    else:
      self.px_ssh.sendline(str(self.send_commandline))  
      self.print_console(self,0)     
      self.snd_return_code_call(self,self.q_s,'pxssh')    
        
  def can_referenceboard(self, sel_snd,q_s,q_r,can): 
     
    self.sel_snd= sel_snd
    self.q_s= q_s
    self.q_r= q_r
    self.can= can
    
    if self.can in config_parser.interfacecanlistref:
      indx=config_parser.interfacecanlistref.index(self.can)
      if config_parser.interfacecanlistref[indx+1].find(",") >=0:
        self.type_list=config_parser.interfacecanlistref[indx+1].split(",")   
        self.send_commandline=self.type_list[0]        
        self.rcv_commandline=self.type_list[1]
        print("Send commandline: {}".format(self.send_commandline))
        print("Receive commandline: {}".format(self.rcv_commandline)) 
                    
    if self.sel_snd=='snd':
      self.px_ssh_1.sendline(self.send_commandline)  
      self.print_console(self,1)
      self.snd_return_code_call(self,self.q_s,'pxssh1')      
    else:
      self.px_ssh_1.sendline(str(self.rcv_commandline))
      self.print_console(self,1)
      self.rcv_return_code_call(self,self.q_r,'pxssh1')  
    
  def return_code_call(self):
    
      self.px_ssh.sendline('echo $?')
      self.print_console(self,0)   
      self.ret_code= '\r\n'.join(str(self.px_ssh.before.splitlines()[1:]))
      self.ret_code=self.ret_code[9]
      return self.ret_code
  '''------------------------------------------------------------------------------------------------------------------------------------------------------------------
                  Class method pxssh_prompt(self,order,commandline ): executes the executable file on the board 
                  Input parameters:
                  self           --> Instance of a class plug FS_BoardPlug
                  order          --> interface type name(i2c_a, spi_a, etc.....)
                  commandline    --> commandline for executing the tests(./rcv_uart /dev/ttymxc3, etc ........)
  --------------------------------------------------------------------------------------------------------------------------------------------------------------------'''                     
  def pxssh_prompt(self,order,commandline,frequency):

    self.order=order
    self.commandline=commandline
    self.frequency=frequency

    if(self.order == 'I2C_A' or
       self.order == 'I2C_B' or
       self.order == 'I2C_C' or
       self.order == 'I2C_D'):
         
       self.ssh=paramiko.SSHClient()
       self.ssh.load_system_host_keys()
       self.ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
       self.ssh.connect(FS_BoardPlug.dut_id_test_board, 22, "root", "root")
       
       if(self.frequency==50):
        new_freq=input('FREQUENCY CHANGED TO 50 KHZ?: (yes/no) \n')
       if(self.frequency==100):
        new_freq=input('FREQUENCY CHANGED TO 100 KHZ?: (yes/no) \n')
       if(self.frequency==150):
        new_freq=input('FREQUENCY CHANGED TO 400 KHZ?: (yes/no) \n')
       
       while True:
        if new_freq =='yes':
          break
        elif new_freq == 'no':
          print("Rebooting.................Timeout: 10 min")
          stdin, stdout, stderr=self.ssh.exec_command('/sbin/reboot -f > /dev/null 2>&1 &')
          input('Type anything if rebooted and press ENTER!')
          break
        else:
          new_freq=input('FREQUENCY CHANGED?: (yes/no) \n')

       FS_BoardPlug.px_ssh=pxssh.pxssh()
       FS_BoardPlug.px_ssh.login(FS_BoardPlug.dut_id_test_board,FS_BoardPlug.username,FS_BoardPlug.password) 
       print("SSH connection established \n")
       self.px_ssh.sendline(commandline)
       self.print_console(self,0)   
       return self.return_code_call(self)
         
    #hdmi alongwith one blank space and lcd/ldb alongwith two blank space check config_ini file  inorder to 
    #get desired names of Phases 
    elif(self.order == 'HDMI 'or
         self.order == 'LCD  ' or
         self.order == 'LDB  ' or
         self.order == 'Touchscreen'):  
         
         self.ssh=paramiko.SSHClient()
         self.ssh.load_system_host_keys()
         self.ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
         self.ssh.connect(FS_BoardPlug.dut_id_test_board, 22, "root", "root")
         
         if self.order == 'Touchscreen':
         
          touch_status_RGB = input('Status of'+" "+ self.order+" "+' lcd ? (ok/nok) \n')
          touch_status_LDB = input('Status of'+" "+ self.order+" "+' ldb ? (ok/nok) \n')
          touch_status_HDMI = input('Status of'+" "+ self.order+" "+' hdmi ? (ok/nok) \n')
          
          while True:
            if ( touch_status_RGB == 'ok' and 
               touch_status_LDB == 'ok' and
               touch_status_HDMI =='ok'):
               
               return 0
               break 
            elif ( touch_status_RGB == 'nok' or 
               touch_status_LDB == 'nok' or
               touch_status_HDMI =='nok'):
               
               return 1
               break
            else:
               print("Provide the proper touch status of lcd/ldb/hdmi: (ok/nok) \n")
               touch_status_RGB = input('Status of'+" "+ self.order+" "+' lcd ? (ok/nok) \n')
               touch_status_LDB = input('Status of'+" "+ self.order+" "+' ldb ? (ok/nok) \n')
               touch_status_HDMI = input('Status of'+" "+ self.order+" "+' hdmi ? (ok/nok) \n')
          
         else:
          display_typ=input('Configured'+" "+self.order +'in device tree?: (y/n) \n')
         
          while True:
            if display_typ =='y':
              status=input('Status of'+" "+ self.order+" "+'? (ok/nok) \n')
              while True:
                if status=='ok':
                  return 0
                  break
                elif status=='nok':
                  return 1
                  break
                else:
                  print("Provide the proper status of display: (ok/nok) \n")
                  status=input('Status of'+" "+ self.order+" "+'? (ok/nok) \n')
              break
            
            elif display_typ == 'n':
              print("Rebooting.................")
              stdin, stdout, stderr=self.ssh.exec_command('/sbin/reboot -f > /dev/null 2>&1 &')
              input('Type anything if the board is rebooted and network is configured using UDHCP !')
            
              status=input('Status of'+" "+ self.order+" "+'? (ok/nok) \n')
              while True:
                if status=='ok':
                  FS_BoardPlug.px_ssh=pxssh.pxssh()
                  FS_BoardPlug.px_ssh.login(FS_BoardPlug.dut_id_test_board,FS_BoardPlug.username,FS_BoardPlug.password) 
                  print("SSH connection established \n")
                  return 0
                  break
                elif status=='nok':
                  FS_BoardPlug.px_ssh=pxssh.pxssh()
                  FS_BoardPlug.px_ssh.login(FS_BoardPlug.dut_id_test_board,FS_BoardPlug.username,FS_BoardPlug.password) 
                  print("SSH connection established \n")
                  return 1
                  break
                else:
                  print("Provide the proper status of display: (ok/nok) \n")
                  status=input('Status of'+" "+ self.order+" "+'? (ok/nok) \n')
              break
          
            else:
              display_typ=input('Configured'+" "+self.order +'in device tree?: (y/n) \n')
            
         FS_BoardPlug.px_ssh=pxssh.pxssh()
         FS_BoardPlug.px_ssh.login(FS_BoardPlug.dut_id_test_board,FS_BoardPlug.username,FS_BoardPlug.password) 
         print("SSH connection established \n")
         
    elif(self.order =='I2S' or self.order == 'I2C_C'):
        
         self.px_ssh.sendline(commandline)
         self.print_console(self,0)
         
         status=input('Status of Audio?: (ok/nok) \n')
         while True:
              if status=='ok':
                return 0
                break
              elif status=='nok':
                return 1
                break
              else:
                status=input('Status of Audio?: (ok/nok) \n')
                
    elif(self.order=='WLAN'):
    
        dummy_ret= -1
        return dummy_ret

    elif(self.order == 'SPI_A' or
         self.order == 'SPI_B' or
         self.order == 'SPI_C' or
         self.order == 'SPI_D'):  
         
         self.px_ssh.sendline(commandline)
         self.print_console(self,0)      
         return self.return_code_call(self)
         
    elif(self.order == 'Internes Flash'):  
         
         self.px_ssh.sendline(commandline)
         self.print_console(self,0)      
         return self.return_code_call(self)
         
    elif(self.order == 'RTC'):
      
         loc=locale.setlocale(locale.LC_ALL,'')
         date = datetime.datetime.today().strftime(u'%Y-%m-%d')
         time =  datetime.datetime.today().strftime(u'%H:%M:%S')
         commandline= commandline + " " +"'"+date+" "+time+"'"+" "+ "&& " + "hwclock -w"
         self.px_ssh.sendline(commandline)
         self.print_console(self,0)
         self.ret_second= '\r\n'.join(str(self.px_ssh.before.splitlines()[1:]))    
         self.ret_second_bfr=self.ret_second[60]+self.ret_second[63]
         self.ret_year_bfr= self.ret_second[81]+self.ret_second[84]+self.ret_second[87]+self.ret_second[90]
         
         import time
         time.sleep(10)
         
         self.px_ssh.sendline('hwclock')
         self.print_console(self,0)
         self.ret_second= '\r\n'.join(str(self.px_ssh.before.splitlines()[1:]))
         self.ret_second_aftr=self.ret_second[60]+self.ret_second[63]
         self.ret_year_aftr= self.ret_second[69]+self.ret_second[72]+self.ret_second[75]+self.ret_second[78]
         
         if (int(self.ret_second_bfr) >= 50 and int(self.ret_second_aftr) < 61):
          self.return_value=60+int(self.ret_second_aftr)-int(self.ret_second_bfr)
         else:
          self.return_value=int(self.ret_second_aftr)-int(self.ret_second_bfr) 
          
         self.return_year_value=int(self.ret_year_aftr)-int(self.ret_year_bfr)
         print("Year before: {}\nYear after: {}\nSecond difference: {}\n".format(self.ret_year_bfr,self.ret_year_aftr,self.return_value))
         
         if self.return_year_value==0:
          if (self.return_value == 9 or
             self.return_value == 10 or
             self.return_value == 11):
            return 0
          else:
            return 1
         else:
            return 1

    elif(self.order == 'SD_A' or
         self.order == 'SD_B' or
         self.order == 'USB_H' or
         self.order == 'USB_D' or
         self.order == 'sda3' or
         self.order == 'eMMC'):  
         
         if(self.order=='USB_D'):
          usb_device_status=input(commandline+" "+'? (ok/nok) \n')
          while True:
              if usb_device_status=='ok':
                return 0
                break
              elif usb_device_status=='nok':
                return 1
                break
              else:
                usb_device_status=input(commandline+" "+'? (ok/nok) \n')
          
         else:
          self.px_ssh.sendline(commandline)
          self.print_console(self,0)      
          self.ret_1=self.return_code_call(self) 

          self.px_ssh.sendline('cd /root/ && umount /mnt/')
          self.print_console(self,0)
          self.ret_2=self.return_code_call(self) 
         
         if (self.ret_1 == '0' and self.ret_2 == '0'):
          return 0
         else:
          return 1
          
    
    #After pwm there is always two blank space like lcd and ldb     
    elif(self.order == 'PWM_A' or
         self.order == 'PWM_B' or
         self.order == 'PWM_C' or
         self.order == 'PWM  (BL_CTRL)'):   
         
         if (self.order == 'PWM  (BL_CTRL)'):
          backlight_status=input( commandline+" "+'? (ok/nok) \n')
          while True:
              if backlight_status=='ok':
                return 0
                break
              elif backlight_status=='nok':
                return 1
                break
              else:
                print("Provide the proper status of backlight control: (ok/nok) \n")
                backlight_status=input('Status of'+" "+ self.order+" "+'(ok/nok) \n')
                 
         if(self.order =='PWM_B'):
          return -1
          
         else:
          self.px_ssh.sendline(commandline)
          self.print_console(self,0)    
          return self.return_code_call(self)
         
    elif(self.order == 'ETH_A' or
         self.order == 'ETH_B' or
         self.order == 'ETH_C' or
         self.order == 'ETH_D'):
            
         #ni.ifaddresses('eth0')
         #ip = ni.ifaddresses('eth0')[ni.AF_INET][0]['addr']
         #commandline=commandline+" "+ip
         if self.order=='ETH_B':
          return -1
          
         status=input( commandline+" "+'? (ok/nok) \n')
         while True:
              if status=='ok':
                return 0
                break
              elif status=='nok':
                return 1
                break
              else:
                print("Provide the proper status of ethernet: (ok/nok) \n")
                status=input('Status of'+" "+ self.order+" "+'(ok/nok) \n')    

    elif(self.order == 'UART_A' or
         self.order == 'UART_B' or
         self.order == 'UART_C' or
         self.order == 'UART_D'):
         
         import time       

         q_s=queue.Queue()
         q_r=queue.Queue()
         
         t1=threading.Thread(target=self.uart_referenceboard,args=(self,'send',q_s,q_r,self.order))
         t2=threading.Thread(target=self.testingboard,args=(self,'rcv',q_s,q_r,self.commandline))

         t2.start()
         time.sleep(2)
         t1.start()
         t2.join()
         t1.join()

         self.ret_send=int(q_s.get())
         print('Thread1',self.ret_send)
         self.ret_rcv=int(q_r.get())
         print('Thread2',self.ret_rcv)
          
         t3=threading.Thread(target=self.uart_referenceboard,args=(self,'rcv',q_s,q_r,self.order))
         t4=threading.Thread(target=self.testingboard,args=(self,'send',q_s,q_r,self.commandline))

         t3.start()
         time.sleep(2)        
         t4.start()
         t3.join()
         t4.join()
         
         self.ret_send_1=int(q_s.get())
         print('Thread3',self.ret_send_1)
         self.ret_rcv_1=int(q_r.get())
         print('Thread4',self.ret_rcv_1)
        
         if (self.ret_send==0 and self.ret_rcv==0 and
            self.ret_send_1==0 and self.ret_rcv_1==0):
          return 0
         else:
          return 1   
             
    elif(self.order=='CAN_A' or
        self.order =='CAN_B' or
        self.order =='CAN_C' or
        self.order =='CAN_D'):
        
        print(self.commandline)
        import time
        
        q_s=queue.Queue()
        q_r=queue.Queue()
       
        t1=threading.Thread(target=self.testingboard,args=(self,'rcv',q_s,q_r,self.commandline))
        t2=threading.Thread(target=self.can_referenceboard,args=(self,'snd',q_s,q_r,self.order))
        
        t1.start()
        time.sleep(2)
        t2.start()
        t1.join()
        t2.join()

        self.ret_rcv_can=int(q_r.get())
        print('Thread2',self.ret_rcv_can)
        self.ret_send_can=int(q_s.get())
        print('Thread1',self.ret_send_can)
        
        t3=threading.Thread(target=self.testingboard,args=(self,'rcv',q_s,q_r,self.commandline))
        t4=threading.Thread(target=self.can_referenceboard,args=(self,'snd',q_s,q_r,self.order))
        t3.start()
        time.sleep(2)
        t4.start()
        t4.join()
        t3.join()
        self.ret_rcv_can_1=int(q_r.get())
        print('Thread2',self.ret_rcv_can_1)
        self.ret_send_can_1=int(q_s.get())
        print('Thread1',self.ret_send_can_1)
        
        if (self.ret_send_can==0 and self.ret_rcv_can==0 and
            self.ret_rcv_can_1==0 and self.ret_send_can_1==0):
            
          return 0
        else:
          return 1
    else:
        pass

        

'''------------------------------------------------------------------------------------------------------------------------------------------------------------------
                  Method select_interface_to_test(select_board,select_interface): select an interface or multiple interfaces of a particular board
                  Input parameters:
                  select_interface --> Interface to be selected for a test
                  select_board     --> Board to undergo a test
----------------------------------------------------------------------------------------------------------------------------------------------------------------------'''
def select_interface_to_test(select_board,select_interface):
      
      if (select_interface=='UART_A' or select_interface=='UART_B' or
          select_interface=='UART_C' or select_interface=='UART_D'):

               test_phases_uart=[
                        test_UART.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
               test=htf.Test(test_phases_uart)
               test.add_output_callbacks(server.publish_final_state)
               test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
               
      elif (select_interface=='all0'):              

               uart_type_list=config_parser.INTERFACE_ALL
               test_phases_uart=[
                        test_UART.with_args(interface=uart_type)
                        for uart_type in uart_type_list
                        ]              
               test=htf.Test(test_phases_uart)
               test.add_output_callbacks(server.publish_final_state)
               test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                 
      elif (select_interface=='SPI_A' or select_interface=='SPI_B' or
            select_interface=='SPI_C' or select_interface=='SPI_D'):

                test_phases_spi=[
                        test_SPI.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_spi)
                test.add_output_callbacks(server.publish_final_state)            
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all1'):

                spi_type_list=config_parser.INTERFACE_ALL
                test_phases_spi=[
                        test_SPI.with_args(interface=spi_type)
                        for spi_type in spi_type_list
                        ]              
                test=htf.Test(test_phases_spi)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
            
      elif (select_interface=='I2C_A' or select_interface=='I2C_B' or
            select_interface=='I2C_C' or select_interface=='I2C_D'):

                test_phases_i2c=[
                        test_I2C.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_i2c)
                test.add_output_callbacks(server.publish_final_state)          
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif(select_interface=='all2'): 

               i2c_type_list=config_parser.INTERFACE_ALL
               test_phases_i2c=[
                        test_I2C.with_args(interface=i2c_type)
                        for i2c_type in i2c_type_list
                        ]              
               test=htf.Test(test_phases_i2c)
               test.add_output_callbacks(server.publish_final_state)
               test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='CAN_A' or select_interface=='CAN_B' or
            select_interface=='CAN_C' or select_interface=='CAN_D'):
            
                test_phases_can=[
                        test_CAN.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_can)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all3'):

                can_type_list=config_parser.INTERFACE_ALL
                test_phases_can=[
                        test_CAN.with_args(interface=can_type)
                        for can_type in can_type_list
                        ]              
                test=htf.Test(test_phases_can)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
        
      elif (select_interface=='PWM_A' or select_interface=='PWM_B' or
            select_interface=='PWM  (BL_CTRL)'):
      #elif (select_interface=='pwm  '):
                test_phases_pwm=[
                        test_PWM.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_pwm)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all4'):
                pwm_type_list=config_parser.INTERFACE_ALL
                test_phases_pwm=[
                        test_PWM.with_args(interface=pwm_type)
                        for pwm_type in pwm_type_list
                        ]              
                test=htf.Test(test_phases_pwm)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))  
                
      elif (select_interface=='ETH_A' or select_interface=='ETH_B'):

                test_phases_network=[
                        test_NETWORK.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_network)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all5'):
                network_type_list=config_parser.INTERFACE_ALL
                test_phases_network=[
                        test_NETWORK.with_args(interface=network_type)
                        for network_type in network_type_list
                        ]              
                test=htf.Test(test_phases_network)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))      
                
      elif (select_interface=='SD_A' or select_interface=='SD_B'):

                test_phases_sdcard=[
                        test_SDCARD.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_sdcard)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all6'):
                sdcard_type_list=config_parser.INTERFACE_ALL
                test_phases_sdcard=[
                        test_SDCARD.with_args(interface=sdcard_type)
                        for sdcard_type in sdcard_type_list
                        ]              
                test=htf.Test(test_phases_sdcard)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='RTC'):

                test_phases_rtc=[
                        test_RTC.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_rtc)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all7'):
                rtc_type_list=config_parser.INTERFACE_ALL
                test_phases_rtc=[
                        test_RTC.with_args(interface=rtc_type)
                        for rtc_type in rtc_type_list
                        ]              
                test=htf.Test(test_phases_rtc)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
      #hdmi with one space and lcd along with two white spaces          
      elif (select_interface=='HDMI ' or select_interface=='LCD  'or
            select_interface=='Touchscreen' or select_interface=='LDB  '):
                test_phases_display=[
                        test_DISPLAY.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_display)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all9'):
                display_type_list=config_parser.INTERFACE_ALL
                test_phases_display=[
                        test_DISPLAY.with_args(interface=display_type)
                        for display_type in display_type_list
                        ]              
                test=htf.Test(test_phases_display)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='eMMC'):

                test_phases_emmc=[
                        test_EMMC.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_emmc)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all10'):
                emmc_type_list=config_parser.INTERFACE_ALL
                test_phases_emmc=[
                        test_EMMC.with_args(interface=emmc_type)
                        for emmc_type in emmc_type_list
                        ]              
                test=htf.Test(test_phases_emmc)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))  
                
      elif (select_interface=='Internes Flash'):

                test_phases_flash=[
                        test_FLASH.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_flash)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all11'):
                flash_type_list=config_parser.INTERFACE_ALL
                test_phases_flash=[
                        test_FLASH.with_args(interface=flash_type)
                        for flash_type in flash_type_list
                        ]              
                test=htf.Test(test_phases_flash)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='I2S' or select_interface=='I2C_C'):

                test_phases_i2s=[
                        test_I2S.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_i2s)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all13'):
                i2s_type_list=config_parser.INTERFACE_ALL
                test_phases_i2s=[
                        test_I2S.with_args(interface=i2s_type)
                        for i2s_type in i2s_type_list
                        ]              
                test=htf.Test(test_phases_i2s)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                  
      elif (select_interface=='usb_h'or select_interface=='usb_d'):

                test_phases_usb=[
                        test_USB.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_usb)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all8'):
                usb_type_list=config_parser.INTERFACE_ALL
                test_phases_usb=[
                        test_USB.with_args(interface=usb_type)
                        for usb_type in usb_type_list
                        ]              
                test=htf.Test(test_phases_usb)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))  
                
      elif (select_interface=='WLAN'):

                test_phases_wlan=[
                        test_WLAN.with_args(interface=config_parser.INTERFACE_A_B_LIST)
                        ]
                test=htf.Test(test_phases_wlan)
                test.add_output_callbacks(server.publish_final_state)        
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))
                
      elif (select_interface=='all12'):
                wlan_type_list=config_parser.INTERFACE_ALL
                test_phases_wlan=[
                        test_WLAN.with_args(interface=wlan_type)
                        for wlan_type in wlan_type_list
                        ]              
                test=htf.Test(test_phases_wlan)
                test.add_output_callbacks(server.publish_final_state)
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))                   
      else:
                test_phases=(
                        [                       
                        test_I2C.with_args(interface=i2c_type)
                        for i2c_type in config_parser.i2clist
                        ]
                        ,
                        [
                        test_SPI.with_args(interface=spi_type)
                        for spi_type in config_parser.spilist
                        ]
                        ,
                        [
                        test_UART.with_args(interface=uart_type)
                        for uart_type in config_parser.uartlist
                        ]
                        ,
                        [
                        test_CAN.with_args(interface=can_type)
                        for can_type in config_parser.canlist
                        ]
                        ,
                        [
                        test_PWM.with_args(interface=pwm_type)
                        for pwm_type in config_parser.pwmlist
                        ]
                        ,
                        [
                        test_NETWORK.with_args(interface=network_type)
                        for network_type in config_parser.networklist
                        ]
                        ,
                        [
                        test_SDCARD.with_args(interface=sdcard_type)
                        for sdcard_type in config_parser.sdcardlist
                        ]
                        ,                        
                        [
                        test_RTC.with_args(interface=rtc_type)
                        for rtc_type in config_parser.rtclist
                        ]
                        ,
                        [
                        test_USB.with_args(interface=usb_type)
                        for usb_type in config_parser.usblist
                        ]
                        ,
                        [
                        test_DISPLAY.with_args(interface=display_type)
                        for display_type in config_parser.displaylist
                        ]
                        ,
                        [
                        test_EMMC.with_args(interface=emmc_type)
                        for emmc_type in config_parser.emmclist
                        ]
                        ,
                        [
                        test_FLASH.with_args(interface=flash_type)
                        for flash_type in config_parser.flashlist
                        ]
                        ,
                        [
                        test_WLAN.with_args(interface=wlan_type)
                        for wlan_type in config_parser.wlanlist
                        ]
                        ,
                        
                        [
                        test_I2S.with_args(interface=i2s_type)
                        for i2s_type in config_parser.i2slist
                        ]
                        ,                        
                        )
                test=htf.Test(*test_phases)
                test.add_output_callbacks(server.publish_final_state)  
                test.execute(test_start=FS_BoardPlug.pxssh_login(FS_BoardPlug))

#creating a header in record table and outputting in excel

def create_excel_table():

  global workbook

  sheet = workbook.active  
  header_value=('referenceboard', 'date/time-stamp', 'testedboard','interface', 'comment', 'revision', 'author', 'status')
  cell_range=sheet['A1':'H1']
  counter=0
  row_a=sheet[1]
  sheet.auto_filter.ref="A1:H1"
  # Creating a few styles in the cell
  header=NamedStyle(name="header")
  header.font=Font(bold=True,size=15)
  header.border=Border(bottom=Side(border_style="thick"))
  header.alignment=Alignment(horizontal="center", vertical="center")

  for cell in row_a:
      cell.style=header
      cell.value=header_value[counter]
      counter=counter+1
  workbook.save(filename="validtablelist.xlsx")

def valid_table_excel(interface,transient_flag,comment): 

  global workbook

  sheet = workbook.active
  list_value=['efusa9', '2020-03-05','PicocoreSoloX','I2C','Baudrate:9600','Frontend_03_20_1.0','Ravindra','OK']
  list_value[6]=FS_BoardPlug.name
  #date and time record
  loc=locale.setlocale(locale.LC_ALL,'')
  localized_time = datetime.datetime.today().strftime(u'%d.%m.%Y %H:%M:%S')
  list_value[1]=localized_time
  list_value[2]=config_parser.BOARD_NAME
  list_value[3]=interface
  list_value[4]=comment
  cell_range=sheet['A1':'H1'] 
  
  global sel_row

  #Data entry to a table
  if transient_flag==0:
    list_value[7]='OK'
    sel_row=sel_row+1
    row_a= sheet[sel_row]    
    counter=0
    for cell in row_a:
      cell.value=list_value[counter]
      if cell.value=='OK':
            cell.font=Font(bold=True,color=colors.GREEN)          
      counter=counter+1
    #Data entry to a table
  elif transient_flag==1:
    list_value[7]='NOK'
    sel_row=sel_row+1
    row_a= sheet[sel_row]    
    counter=0
    for cell in row_a:
      cell.value=list_value[counter]
      if cell.value=='NOK':
        cell.font=Font(bold=True,color=colors.RED)            
      counter=counter+1 
    #Data entry to a table 
  else:  
    list_value[7]='NA'
    sel_row=sel_row+1
    row_a= sheet[sel_row]    
    counter=0
    for cell in row_a:
      cell.value=list_value[counter]
      if cell.value=='NA':
        cell.font=Font(bold=True,color=colors.BLUE)            
      counter=counter+1 
      
  workbook.save(filename="validtablelist.xlsx")

def heading_contents_table(col_names,tx_cells,bold):

  for idx, name in enumerate(col_names):
    paragraph=tx_cells[idx].paragraphs[0]
    run=paragraph.add_run(name)
    run.bold=True if bold==0 else False
    
def test_report_template():
  
  global doc, recording_table
  doc.add_heading('Spezifizierung',2)

  para=doc.add_paragraph('Die folgende Liste enthält Angaben welche Peripherie Einheiten bei der Boardfamilie ' )
  para.add_run('efus').bold=True
  para.add_run('TM ').font.superscript=True
  para.add_run('zu testen sind. Die Spezifizierungen zu den einzelnen Tests sind im Dokument QMAA Schnittstellentest(ID:2685), Revision: 011 zu finden.')
  para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

  col_names=('Fertigungsbaugruppe','Revision','efus-SINTF. Revision','OS','Bootloader Version','Kernel Version','Filename')
  table_main=doc.add_table(rows=2,cols=len(col_names))
  table_main.allow_autofit=True
  table_main.style="Light Grid"
  table_main.rows[0].style="background-color:gray"

  tx_cells=table_main.rows[0].cells
  heading_contents_table(col_names,tx_cells,0)
  
  col_names=['EFUS-A9','1.21','1.30','Linux','NBoot VN41 U-Boot2018.03','4.9.88','fsimx6-B2020.04']
  tx_cells=table_main.rows[1].cells
  heading_contents_table(col_names,tx_cells,1)

  space_table=doc.add_table(rows=1,cols=1)
  space_table.style="Light Grid"
  tx_cells=space_table.rows[0].cells
  tx_cells[0].text=''
 
  recording_table= doc.add_table(rows=30,cols=7)
  recording_table.allow_autofit=True
  recording_table.style="Light Grid"
    
  col_names=['Peripherie\u2071','ID\u2071\u2071','Ergebnis OS\u2071\u2071\u2071','Ergebnis UBoot\u2071\u1d5b','Kommentar','MA','Datum']
  tx_cells=recording_table.rows[0].cells
  heading_contents_table(col_names,tx_cells,0)
  
  end_para=doc.add_paragraph("_____________________________________________")
  end_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  end_para_symbl_i=doc.add_paragraph('\u2071Bezeichnung der zu testenden Peripherieeinheit.')
  end_para_symbl_i.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  end_para_symbl_ii=doc.add_paragraph('\u2071\u2071Kennung des durchzuführenden Tests. Kennung muss identisch zur Anforderung sein. Die Anforderung ist im Dokument „SW EW Checklist“ mit gleicher ID erfasst. Nummerierung \
                    nicht ändern sonst stimmt Korrelation mit Anforderungen nicht überein.')
  end_para_symbl_ii.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  end_para_symbl_iii=doc.add_paragraph('\u2071\u2071\u2071Hier eintragen: OK = Test bestanden, NOK = Test nicht bestanden und NA = Test nicht zutreffend. Bitte bei NOK und NA im Feld Kommentar ergänzen warum Test fehlschlägt \
                     oder warum Test nicht zutreffend ist.')
  end_para_symbl_iii.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
  end_para_symbl_iv=doc.add_paragraph('\u2071\u1d5bHier eintragen: OK = Test bestanden, NOK = Test nicht bestanden und NA = Test nicht zutreffend. Bitte bei NOK und NA im Feld Kommentar ergänzen warum Test fehlschlägt \
                     oder warum Test nicht zutreffend ist.')
  end_para_symbl_iv.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY      

  doc.save('Peripherie Test efusA9 Linux.docx')

def test_report(interface, transient_flag,comment,id_value):

 global doc, recording_table,count
 print(id_value)
 loc=locale.setlocale(locale.LC_ALL,'')
 localized_time = datetime.datetime.today().strftime(u'%d.%m.%Y %H:%M:%S')
 tx_cells=recording_table.rows[int(id_value)].cells
 
 tx_cells[0].text=interface
 tx_cells[1].text=str(id_value)
 if (int(id_value) == 4 or
    int(id_value) == 11 or
    int(id_value) == 12):
    tx_cells[3].text='OK'
 else:
    tx_cells[3].text='NA'    
 tx_cells[4].text=comment
 tx_cells[5].text=FS_BoardPlug.name
 tx_cells[6].text=localized_time
 if transient_flag==0:
  tx_cells[2].text='OK'
  
 elif transient_flag==1:
  tx_cells[2].text='NOK'
  
 else:
  tx_cells[2].text='NA'
 
 count=count+1
 doc.save('Peripherie Test efusA9 Linux.docx')
 
# Note: phase name and total_time measurement use {} formatting with args
# passed into the phase so each phase has a unique name.
@htf.PhaseOptions(name='I2C_{interface[4]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
) 
def test_I2C(test,interface):

  test.logger.info('------I2C test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('--------------------------I2C test starts---------------------------------')
  #Add command for removing the files if it existed already
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'i2c')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'i2c')
  os.remove('/home/developer/Desktop/selftest_i2c')
  return_code_50=int(FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,50))
  return_code_100=int(FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,100))
  return_code_150=int(FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,150))
  if (return_code_50 ==0 and return_code_100==0 and return_code_150==0):
   return_code=0
  else:
    return_code=1
  end=time.time()
  transient_flag=return_code
  comment="testing i2c-device with address,"\
          "\nblinking LEDs,\n"\
          "with different frequencies."
  valid_table_excel(interface,transient_flag, comment)       
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  test_report(interface, transient_flag,comment,id_value)
  info=test.logger.info('--------I2C test phase ends--------------')

@htf.PhaseOptions(name='SPI_{interface[4]}')
@htf.plug(FS_BoardPlug)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
) 
def test_SPI(test,interface):

  test.logger.info('---------SPI test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------SPI test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'spi')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'spi')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/selftest_spi')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="max speed: 1000 KHz(1 MHz)\n"\
          "max speed: 10000 KHz(10 MHz)\n"\
          "max speed: 24000 KHz(24 MHz)"
  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------SPI test phase ends------')
  
@htf.PhaseOptions(name='UART_{interface[5]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
) 
def test_UART(test,interface):

  test.logger.info('------UART test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          cal_id=command
          if cal_id.find("*")>=0:
           command_id=cal_id.split("*")
           id_value=command_id[1]
           id_value=id_value.split(",")
           id_value=id_value[0]
  else:
          pass
  start = time.time()
  print('-------------------UART test starts-------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'send_uart_abr')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'send_uart_rs485_abr')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'rcv_uart_abr')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'rcv_uart_rs485_abr')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'uart')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/send_uart_rs485_abr')
  os.remove('/home/developer/Desktop/rcv_uart_rs485_abr')
  os.remove('/home/developer/Desktop/send_uart_abr')
  os.remove('/home/developer/Desktop/rcv_uart_abr')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment=  "Tested with:\n"\
            "Baud rate: 9600\n"\
            "Baud rate: 38400\n"\
            "Baud rate: 115200\n"
  valid_table_excel(interface,transient_flag, comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------UART test phase ends------')

@htf.PhaseOptions(name='CAN_{interface[4]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)
def test_CAN(test,interface):

  test.logger.info('-------CAN test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          cal_id=command
          if cal_id.find("*")>=0:
           command_id=cal_id.split("*")
           id_value=command_id[1]
           id_value=id_value.split(",")
           id_value=id_value[0]
  else:
          pass
  start = time.time()
  print('---------------CAN test starts--------------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'snd_can_afs')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'rcv_can_afs')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'can')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/snd_can_afs')
  os.remove('/home/developer/Desktop/rcv_can_afs')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment=  "Tested with:\n"\
            "Baud rate: 100000\n"\
            "Baud rate: 120000\n"\
            "Baud rate: 1000000"
  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------CAN test phase ends------') 

@htf.PhaseOptions(name='PWM {interface[3]}{interface[4]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int),
    htf.Measurement('transient_flag_1').equals(-1,type=int)
)  
def test_PWM(test,interface):

  test.logger.info('---------PWM test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------PWM test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'pwm')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'pwm')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/selftest_pwm_new')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  comment="period is set to: 500\n"\
          "period is set to: 1000."\

  if interface =='PWM_B':
    test.measurements.transient_flag_1 = return_code
    test.measurements.return_code = 0
    transient_flag_1=return_code
    valid_table_excel(interface,transient_flag_1,comment)
    test_report(interface, transient_flag_1,comment,id_value)
    
  elif interface == 'PWM  (BL_CTRL)':
    test.measurements.return_code = return_code
    test.measurements.transient_flag_1 = -1
    transient_flag=return_code
    valid_table_excel(interface,transient_flag,comment)
    test_report('PWM (BL_CTRL)', transient_flag,'Manual test',id_value)
  else:
    test.measurements.return_code = return_code
    test.measurements.transient_flag_1 = -1
    transient_flag=return_code
    valid_table_excel(interface,transient_flag,comment)
    test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------PWM test phase ends------')

@htf.PhaseOptions(name='ETH_{interface[4]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int),
    htf.Measurement('transient_flag_1').equals(-1,type=int)
)  
def test_NETWORK(test,interface):

  test.logger.info('---------NETWORK test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------NETWORK test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'network')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'network')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/selftest_network')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  
  comment="files for testing the interfaces\n"\
          "will be transmitted to the\n"\
          "board using SSH."
  if interface =='ETH_B':
    test.measurements.transient_flag_1 = return_code
    test.measurements.return_code = 0
    transient_flag_1=return_code
    valid_table_excel(interface,transient_flag_1,comment)
    test_report(interface, transient_flag_1,comment,id_value)
  else:
    test.measurements.return_code = return_code
    test.measurements.transient_flag_1 = -1
    transient_flag=return_code
    valid_table_excel(interface,transient_flag,comment)
    test_report(interface, transient_flag,comment,id_value)
    
  test.logger.info('------network test phase ends------')

@htf.PhaseOptions(name='SDA_{interface[3]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)  
def test_SDCARD(test,interface):

  test.logger.info('---------SDCARD test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------SDCARD test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'sdcard')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'sdcard')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/selftest_sdcard')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="The USB stick will be connected,\n"\
          "file name test.txt is created,\n"\
          "data is written to the file and\n"\
          "finally file will be read."
  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------sdcard test phase ends------')
  
@htf.PhaseOptions(name='RTC')
@htf.plug(FS_BoardPlug)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)  
def test_RTC(test,interface):

  test.logger.info('---------RTC test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------RTC test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="Date is configured.\n"\
          "Time is configured.\n"\
          "Tested for 10 seconds."
          
  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------RTC test phase ends------')
  
@htf.PhaseOptions(name='{interface[0]}{interface[1]}{interface[2]}{interface[3]}{interface[4]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)  
def test_DISPLAY(test,interface):

  test.logger.info('---------Display test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------Display test starts---------------------------------------')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="Display is configured\n"\
          "according to the test need\n"\
          "in device tree."\
          
  if interface == 'Touchscreen':
    comment = "Tested manually three displays(LCD/LDB/HDMI)."
    valid_table_excel(interface,transient_flag,comment)
    test_report(interface, transient_flag,comment,id_value)
  else:
    valid_table_excel(interface,transient_flag,comment)
    test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------DISPLAY test phase ends------')
  
@htf.PhaseOptions(name='eMMC')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)  
def test_EMMC(test,interface):

  test.logger.info('---------eMMC test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------eMMC test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'sdcard')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'sdcard')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/selftest_sdcard')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="The USB stick will be connected,\n"\
          "file name test.txt is created,\n"\
          "data is written to the file and\n"\
          "finally file will be read."
  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------emmc test phase ends------')
  
@htf.PhaseOptions(name='Internes Flash')
@htf.plug(FS_BoardPlug)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)  
def test_FLASH(test,interface):

  test.logger.info('---------FLASH test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass   
  start = time.time()
  print('------------------FLASH test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'flash')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'flash')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/selftest_flash.sh')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="Test file with 15MB produced\n"\
          "and later on checked the file."\

  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------FLASH test phase ends------')
  
@htf.PhaseOptions(name='Audio+I2C_C')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)  
def test_I2S(test,interface):

  test.logger.info('---------I2S+I2C_C test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------I2S+I2C_C test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  
  comment="Command for playing\n"\
          "right and left speaker was\n"\
          "sent to the board and checked the sound whilst playing."\
          
  comment_i2cc="I2C_C is checked automatically via I2S"
  if return_code==0:       
    valid_table_excel(interface,transient_flag,comment)
    valid_table_excel('I2C_C',transient_flag,comment_i2cc)
    test_report(interface, transient_flag,comment,id_value)
    test_report('I2C_C', transient_flag,comment_i2cc,'9')    
  else: 
    valid_table_excel(interface,transient_flag,comment)
    valid_table_excel('I2C_C',transient_flag,comment_i2cc)
    test_report(interface, transient_flag,comment,id_value)
    test_report('I2C_C', transient_flag,comment_i2cc,'9')
      
  test.logger.info('------I2S+I2C_C test phase ends------')
  
@htf.PhaseOptions(name='USB_{interface[4]}')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(0,type=int)
)  
def test_USB(test,interface):

  test.logger.info('---------USB test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------USB test starts---------------------------------------')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'efusa9')
  FS_BoardPlug.mount_rw(FS_BoardPlug,'testboard')
  FS_BoardPlug.get_exe_file(FS_BoardPlug,'sdcard')
  FS_BoardPlug.send_file_to_remote(FS_BoardPlug,'sdcard')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  os.remove('/home/developer/Desktop/selftest_sdcard')
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="The USB stick will be connected\n"\
          "file name test.txt is created\n"\
          "data is written to the file and\n"\
          "finally file will be read"
  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------USB test phase ends------')
  
@htf.PhaseOptions(name='WLAN')
@htf.plug(FS_BoardPlug)
@htf.TestPhase(timeout_s=10*60)
@htf.measures(
    'total_time',
    htf.Measurement('return_code').equals(-1,type=int)
)  
def test_WLAN(test,interface):

  test.logger.info('---------WLAN test phase starts------')
  if interface.find("=") >= 0:
          type_list=interface.split("=")
          interface=type_list[0]
          command=type_list[1]
          if command.find("*")>=0:
           command_id=command.split("*")
           id_value=command_id[1]
           command=command_id[0]
  else:
          pass
  start = time.time()
  print('------------------Wlan test starts---------------------------------------')
  return_code=FS_BoardPlug.pxssh_prompt(FS_BoardPlug,interface,command,0)
  end=time.time()
  return_code=int(return_code)
  elapsed = time.time() - start
  test.measurements['total_time'] = elapsed
  test.measurements.return_code = return_code
  transient_flag=return_code
  comment="WLAN is not available for efusa9 board.\n"

  valid_table_excel(interface,transient_flag,comment)
  test_report(interface, transient_flag,comment,id_value)
  test.logger.info('------WLAN test phase ends------')

if __name__ == '__main__':

  test = None
  sel_row=1
  count=1
  recording_table= None 
  workbook =Workbook()
  doc=Document()
  
  #clearing the contents of known_hosts file under /home/developer/.ssh/
  os.system('cat /dev/null > /home/developer/.ssh/known_hosts')
  
  create_excel_table()
  test_report_template()
  conf.load(station_server_port='4444')
  
  with station_server.StationServer() as server:
      web_launcher.launch('http://localhost:4444')
      for i in range(1):
      #while(1):
          test=htf.Test()          
          select_interface=str(config_parser.INTERFACE_A_B)
          print("Interface: {}".format(select_interface))
          select_board=str(config_parser.BOARD_NAME)
          print("Board: {}".format(select_board))
          select_interface_ref_uart=str(config_parser.INTERFACE_REF_UART_A)
          print("Reference uart: {}".format(select_interface_ref_uart))
          select_interface_ref_can=str(config_parser.INTERFACE_REF_CAN)
          print("Reference can: {}".format(select_interface_ref_can))
          select_commandline=str(config_parser.INTERFACE_COMMAND)        
          select_interface_to_test(select_board,select_interface)
          #pip3 install -r requirements.txt --user
          

         

          

     

     


